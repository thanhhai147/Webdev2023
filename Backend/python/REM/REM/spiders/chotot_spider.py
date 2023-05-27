from os import link
import scrapy
from twisted.internet import reactor, defer
from scrapy.crawler import CrawlerRunner
from scrapy.utils.log import configure_logging
from scrapy.utils.project import get_project_settings
# from utilities.urlUTL import UrlUTL
from docx import Document
from docx.shared import Inches
import json, uuid, re, io, sys, const

class ChototSpider(scrapy.Spider):
    def __init__(self, doc):
        self.name = "chotot"
        self.doc = doc

    start_urls = [
        f"https://nha.chotot.com/{sys.argv[const.SysArgv.COLLECTION_PATH]}?{sys.argv[const.SysArgv.QUERY_STR]}",
    ]
    
    def start_requests(self):
        for url in self.start_urls:
            yield scrapy.Request(url, callback=self.parse_url, errback=self.err_handler)

    def parse_url(self, response):
        if response.status == "404": raise scrapy.Spider.close("chotot", "Out of page")
        for post in response.css("div.ListAds_ListAds__1z6Pv ul div li.AdItem_wrapperAdItem__1hEwM.AdItem_big__2Sqod"):
            detail_url = post.css("a.AdItem_adItem__2O28x::attr(href)").get()
            detail_url = response.urljoin(detail_url)
            yield scrapy.Request(detail_url, callback=self.parse_detail, errback=self.err_handler)
        # return self.doc.save(f"../../downloads/{self.file_name}.docx")
        # redirect to next page
        # query = UrlUTL.getQuery(response.url)
        # page = int(query["page"])
        # nextpage = response.urljoin(f"?page={page + 1}")
        # yield scrapy.Request(nextpage, callback=self.parse)

    def parse_detail(self, response):
        # get data of the post and convert it to json
        json_string = response.css("#__NEXT_DATA__::text").get()
        json_obj = json.loads(json_string)
        # image
        img_urls = json_obj["props"]["initialState"]["adView"]["adInfo"]["ad"]["images"]
        for url in img_urls:
            self.doc.add_paragraph("")
            para_position = len(self.doc.paragraphs) - 1
            yield scrapy.Request(url, callback= self.parse_image, errback=self.err_handler, cb_kwargs={"para_position": para_position})
        # text
        self.add_text(response, json_obj)
        self.add_spacing()
        self.add_divider()
        self.add_spacing()

    def parse_image(self, response, para_position):
        print()
        with io.BytesIO() as buf:
            buf.write(response.body)
            buf.seek(0)
            self.doc.paragraphs[para_position].add_run().add_picture(buf, width=Inches(3))
    
    def add_text(self, response, json_obj):
        for item in self.get_general_info(json_obj):
            self.doc.add_paragraph(item, style="ListBullet")
        for item in self.get_detail_info(json_obj):
            self.doc.add_paragraph(item, style="ListBullet")
        self.doc.add_paragraph(self.get_link(response), style="Link")
    
    def get_general_info(self, json_obj):
        ad_keys = [
            {"id": "category_name", "label": "Danh mục"}, 
            {"id": "type_name", "label": "Loại hình"}, 
            {"id": "subject", "label": "Tiêu đề"}, 
            {"id": "body", "label": "Miêu tả"},
            {"id": "price_string", "label": "Giá"}
        ]
        for key in ad_keys:
            yield key["label"] + ": " + json_obj["props"]["initialState"]["adView"]["adInfo"]["ad"][key["id"]] + "\n"
    
    def get_detail_info(self, json_obj):
        for parameter in json_obj["props"]["initialState"]["adView"]["adInfo"]["parameters"]:
            if  re.search("ward|area|region", parameter["id"]): continue
            yield parameter["label"] + ": " + parameter["value"] + "\n"
    
    def get_link(self, response):
        return response.url
    
    def add_spacing(self):
        self.doc.add_paragraph("\n" * const.Docx.SPACING_HEIGHT)
    
    def add_divider(self):
        self.doc.add_paragraph(const.Docx.DIVIDER_SYMBOL * const.Docx.DIVIDER_WIDTH)

    # error handler
    def err_handler(self, err):
        self.logger.error(repr(err))

configure_logging()
settings = get_project_settings()
runner = CrawlerRunner(settings)

file_name = uuid.uuid4()

@defer.inlineCallbacks
def crawl():
    doc = Document()
    yield runner.crawl(ChototSpider, doc=doc)
    doc.save(f"downloads/{file_name}.docx")
    reactor.stop()
crawl()
reactor.run()
print(f"{file_name}.docx")